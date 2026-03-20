"""Export module — convert markdown content to PDF or DOCX for download."""
import io
import re
import logging
from typing import Literal

logger = logging.getLogger(__name__)


# ── DOCX export ──────────────────────────────────────────────────────────────
def markdown_to_docx(md_text: str, title: str = "Document") -> bytes:
    """Convert markdown text to a DOCX file and return the bytes."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    for line in md_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("---"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            content = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(p, content)
        else:
            p = doc.add_paragraph()
            _add_formatted_text(p, stripped)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_formatted_text(paragraph, text: str) -> None:
    """Parse simple markdown bold/italic and add runs to a paragraph."""
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


# ── PDF export ───────────────────────────────────────────────────────────────
def markdown_to_pdf(md_text: str, title: str = "Document") -> bytes:
    """Convert markdown text to a PDF file and return the bytes."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable,
    )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        topMargin=0.7 * inch,
        bottomMargin=0.7 * inch,
        leftMargin=0.8 * inch,
        rightMargin=0.8 * inch,
    )

    styles = getSampleStyleSheet()

    styles.add(ParagraphStyle(
        "H1Custom",
        parent=styles["Heading1"],
        fontSize=18,
        textColor=HexColor("#1a1a2e"),
        spaceAfter=8,
    ))
    styles.add(ParagraphStyle(
        "H2Custom",
        parent=styles["Heading2"],
        fontSize=14,
        textColor=HexColor("#16213e"),
        spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        "H3Custom",
        parent=styles["Heading3"],
        fontSize=12,
        textColor=HexColor("#0f3460"),
        spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        "BodyCustom",
        parent=styles["Normal"],
        fontSize=10.5,
        textColor=HexColor("#333333"),
        leading=14,
        spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        "BulletCustom",
        parent=styles["Normal"],
        fontSize=10.5,
        textColor=HexColor("#333333"),
        leading=14,
        leftIndent=20,
        bulletIndent=10,
        spaceAfter=3,
    ))

    story = []

    for line in md_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 6))
            continue

        if stripped.startswith("### "):
            story.append(Paragraph(_md_to_rl(stripped[4:]), styles["H3Custom"]))
        elif stripped.startswith("## "):
            story.append(Paragraph(_md_to_rl(stripped[3:]), styles["H2Custom"]))
        elif stripped.startswith("# "):
            story.append(Paragraph(_md_to_rl(stripped[2:]), styles["H1Custom"]))
        elif stripped.startswith("---"):
            story.append(Spacer(1, 4))
            story.append(HRFlowable(
                width="100%", thickness=0.5,
                color=HexColor("#cccccc"), spaceAfter=4,
            ))
        elif stripped.startswith("- ") or stripped.startswith("* "):
            content = _md_to_rl(stripped[2:])
            story.append(Paragraph(
                f"\u2022  {content}", styles["BulletCustom"]
            ))
        else:
            story.append(Paragraph(_md_to_rl(stripped), styles["BodyCustom"]))

    doc.build(story)
    return buf.getvalue()


def _md_to_rl(text: str) -> str:
    """Convert markdown bold/italic to ReportLab XML tags.

    Escapes XML special characters FIRST, then applies markdown formatting
    so that <b>/<i> tags are not escaped while user text is safe.
    """
    # 1. Escape XML special characters in the raw text
    text = text.replace("&", "&amp;")
    text = text.replace("<", "&lt;")
    text = text.replace(">", "&gt;")
    # 2. Convert markdown bold (**text**) — uses escaped text so safe
    text = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", text)
    # 3. Convert markdown italic (*text*)
    text = re.sub(r"\*(.*?)\*", r"<i>\1</i>", text)
    return text


def export_content(
    md_text: str,
    fmt: Literal["pdf", "docx"] = "pdf",
    title: str = "Document",
) -> bytes:
    """Export markdown content to the requested format."""
    if fmt == "docx":
        return markdown_to_docx(md_text, title)
    return markdown_to_pdf(md_text, title)
